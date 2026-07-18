"""
Document Processor
Handles various document formats (PDF, DOCX, images)
"""
from typing import Dict, Any, List, Optional
from pathlib import Path
import fitz  # PyMuPDF
from PIL import Image
import numpy as np
from docx import Document as DocxDocument
from loguru import logger
import io

from app.config import settings
from app.core.ocr_engine import OCREngineManager


class DocumentProcessor:
    """Process various document formats"""
    
    def __init__(self):
        self.ocr_engine = None
    
    def _ensure_ocr_engine(self):
        """Lazy load OCR engine"""
        if self.ocr_engine is None:
            self.ocr_engine = OCREngineManager.get_engine()
    
    async def process_file(self, file_path: Path) -> Dict[str, Any]:
        """Process file based on type"""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return await self.process_pdf(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return await self.process_image(file_path)
        elif file_extension in ['.docx', '.doc']:
            return await self.process_docx(file_path)
        elif file_extension in ['.txt', '.rtf']:
            return await self.process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    async def process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF file"""
        logger.info(f"Processing PDF: {file_path.name}")
        
        try:
            doc = fitz.open(file_path)
            
            pages_data = []
            all_text = []
            total_confidence = 0.0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Try to extract text directly first
                text = page.get_text()
                
                if text.strip():
                    # Text-based PDF
                    pages_data.append({
                        "page": page_num + 1,
                        "text": text,
                        "confidence": 1.0,
                        "method": "native",
                    })
                    all_text.append(text)
                    total_confidence += 1.0
                else:
                    # Image-based PDF (scanned), use OCR
                    logger.info(f"Page {page_num + 1} requires OCR")
                    
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    img_array = np.array(img)
                    
                    # Perform OCR
                    self._ensure_ocr_engine()
                    ocr_result = self.ocr_engine.extract_text(img_array)
                    
                    pages_data.append({
                        "page": page_num + 1,
                        "text": ocr_result["text"],
                        "confidence": ocr_result["confidence"],
                        "method": "ocr",
                        "blocks": ocr_result.get("blocks", []),
                    })
                    all_text.append(ocr_result["text"])
                    total_confidence += ocr_result["confidence"]
            
            doc.close()
            
            avg_confidence = total_confidence / len(pages_data) if pages_data else 0.0
            
            return {
                "text": "\n\n".join(all_text),
                "pages": pages_data,
                "total_pages": len(pages_data),
                "confidence": avg_confidence,
                "file_type": "pdf",
                "file_name": file_path.name,
            }
        
        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            raise
    
    async def process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image file"""
        logger.info(f"Processing image: {file_path.name}")
        
        try:
            # Load image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Convert to numpy array
            img_array = np.array(image)
            
            # Perform OCR
            self._ensure_ocr_engine()
            ocr_result = self.ocr_engine.extract_text(img_array)
            
            return {
                "text": ocr_result["text"],
                "confidence": ocr_result["confidence"],
                "blocks": ocr_result.get("blocks", []),
                "file_type": "image",
                "file_name": file_path.name,
                "image_size": image.size,
            }
        
        except Exception as e:
            logger.error(f"Failed to process image: {e}")
            raise
    
    async def process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX file"""
        logger.info(f"Processing DOCX: {file_path.name}")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            all_text = paragraphs + tables_text
            
            return {
                "text": "\n".join(all_text),
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "confidence": 1.0,
                "file_type": "docx",
                "file_name": file_path.name,
            }
        
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            raise
    
    async def process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text file"""
        logger.info(f"Processing text file: {file_path.name}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                "text": text,
                "confidence": 1.0,
                "file_type": "text",
                "file_name": file_path.name,
                "length": len(text),
            }
        
        except Exception as e:
            logger.error(f"Failed to process text file: {e}")
            raise
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from document"""
        file_extension = file_path.suffix.lower()
        
        metadata = {
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "file_type": file_extension,
        }
        
        try:
            if file_extension == '.pdf':
                doc = fitz.open(file_path)
                metadata.update({
                    "pages": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "creator": doc.metadata.get("creator", ""),
                    "producer": doc.metadata.get("producer", ""),
                })
                doc.close()
            
            elif file_extension in ['.docx', '.doc']:
                doc = DocxDocument(file_path)
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                })
            
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                image = Image.open(file_path)
                metadata.update({
                    "width": image.width,
                    "height": image.height,
                    "mode": image.mode,
                    "format": image.format,
                })
        
        except Exception as e:
            logger.warning(f"Failed to extract some metadata: {e}")
        
        return metadata
    
    def chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> List[str]:
        """Split text into chunks for processing"""
        chunk_size = chunk_size or settings.CHUNK_SIZE
        chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - chunk_overlap
        
        return chunks
